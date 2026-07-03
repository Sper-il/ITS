# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    print("Reading markdown file...")
    with open('_docx_backup_rewritten.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font for Vietnamese
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith('# ') and not line.startswith('# ' * 2):
            p = doc.add_heading(level=1)
            run = p.add_run(line[2:])
            run.font.size = Pt(16)
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:])
            run.font.size = Pt(14)
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            run = p.add_run(line[4:])
            run.font.size = Pt(12)
        elif line.startswith('**') and line.endswith('**') and line.count('**') == 2:
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('|'):
            # Skip table lines
            continue
        elif line.startswith('```'):
            continue
        elif line.startswith('---'):
            continue
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line, style='List Number')
        else:
            formatted = line
            formatted = re.sub(r'\*\*(.+?)\*\*', r'\1', formatted)
            formatted = re.sub(r'\*(.+?)\*', r'\1', formatted)
            formatted = re.sub(r'`(.+?)`', r'\1', formatted)
            p = doc.add_paragraph(formatted)
            p.paragraph_format.space_after = Pt(6)
    
    doc.save('_docx_backup_rewritten.docx')
    print("Da tao file DOCX thanh cong!")
except ImportError as e:
    print("Loi import: python-docx can be installed with 'pip install python-docx'")
    print(str(e))
except Exception as e:
    print("Loi: " + str(e))
    import traceback
    traceback.print_exc()
